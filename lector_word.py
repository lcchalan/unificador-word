# lector_word.py
from typing import List, Dict, Tuple, Any
import io
import re
import unicodedata

from docx import Document
from docxcompose.composer import Composer
from openpyxl import Workbook


# ---------------------------- Utilidades ----------------------------

def _normalize(text: str) -> str:
    if text is None:
        return ""
    t = unicodedata.normalize("NFKD", text)
    t = "".join(c for c in t if not unicodedata.combining(c))
    t = re.sub(r"\s+", " ", t).strip().casefold()
    return t


def base_title(text: str) -> str:
    return _normalize(text)


def allowed_by_whitelist(level: int, text: str) -> bool:
    # Aquí podrías aplicar tu plantilla de títulos permitidos
    # Por defecto: permitido
    return True


# Bloque = ('h', level, text) | ('p', None, text) | ('t', None, rows:list[list[str]])
Block = Tuple[str, Any, Any]


def _extraer_bloques(doc_bytes: bytes) -> List[Block]:
    """Lee un DOCX y extrae una lista de bloques (headings, párrafos y tablas)."""
    f = io.BytesIO(doc_bytes)
    d = Document(f)
    blocks: List[Block] = []

    # Recorremos elementos en orden: párrafos y tablas
    # python-docx no provee orden directo; iteramos con _body._element
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    for child in d.element.body.iterchildren():
        if isinstance(child, CT_P):
            p = Document().add_paragraph()._p.__class__(child)  # no se usa, solo tip
        # Tomemos el texto y estilo usando un wrapper simple
        if isinstance(child, CT_P):
            para = d._body._element.xpath(".")[0]  # dummy to keep mypy calm
            # Recuperamos el paragraph real:
            paragraph = None
            for p in d.paragraphs:
                if p._p is child:
                    paragraph = p
                    break
            if paragraph is None:
                continue
            style_name = (paragraph.style.name if paragraph.style else "") or ""
            m = re.match(r"Heading\s+(\d+)", style_name, flags=re.I)
            if m:
                level = int(m.group(1))
                text = paragraph.text.strip()
                blocks.append(("h", level, text))
            else:
                text = paragraph.text.strip()
                if text:
                    blocks.append(("p", None, text))

        elif isinstance(child, CT_Tbl):
            # Buscar la tabla que corresponde
            table_obj = None
            for t in d.tables:
                if t._tbl is child:
                    table_obj = t
                    break
            if table_obj is None:
                continue
            rows = []
            for r in table_obj.rows:
                row = []
                for c in r.cells:
                    row.append((c.text or "").strip())
                rows.append(row)
            if rows:
                blocks.append(("t", None, rows))

    return blocks


def headings_from_docx(doc_bytes: bytes) -> List[Dict[str, Any]]:
    """Devuelve [{level:int, text:str}, ...] en orden."""
    bloques = _extraer_bloques(doc_bytes)
    out = []
    for kind, lvl, txt in bloques:
        if kind == "h":
            out.append({"level": int(lvl), "text": str(txt)})
    return out


# ---------------------------- Composición ----------------------------

def _new_doc() -> Document:
    d = Document()
    # DocxComposer necesita un doc no vacío; agregamos un párrafo vacío
    d.add_paragraph("")
    return d


def _append_part(composer: Composer, blocks: List[Block], start: int, end: int):
    """
    Toma los bloques [start:end] (comenzando en un heading) y los envuelve en un doc
    temporal para apendear al composer.
    """
    part = Document()
    # heading original
    kind, lvl, txt = blocks[start]
    h = part.add_heading(level=min(max(int(lvl), 1), 6))
    h.add_run(txt or "")

    for b in blocks[start + 1:end]:
        if b[0] == "p":
            part.add_paragraph(b[2] or "")
        elif b[0] == "t":
            rows = b[2]
            if rows and rows[0]:
                t = part.add_table(rows=len(rows), cols=len(rows[0]))
                for i, row in enumerate(rows):
                    for j, val in enumerate(row):
                        t.cell(i, j).text = val or ""

    composer.append(part)


def _tables_to_xlsx(all_tables: List[Tuple[str, List[List[str]]]]) -> bytes:
    """
    Crea un XLSX con todas las tablas encontradas.
    all_tables = [(fuente, filas), ...]
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Tablas"

    # Cabecera
    ws.append(["Fuente", "Tabla", "Fila", "Columna", "Valor"])

    table_idx = 0
    for source_name, rows in all_tables:
        table_idx += 1
        for r_i, row in enumerate(rows, start=1):
            for c_i, val in enumerate(row, start=1):
                ws.append([source_name, table_idx, r_i, c_i, val or ""])

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio.read()


# ---------------------------- Modos públicos ----------------------------

def procesar(archivos: List[Dict], niveles: List[int], titulos: List[str],
             enforce_whitelist: bool = False) -> Dict[str, bytes]:
    """
    Modo clásico:
      - Un Word 'unificado.docx' con todos los bloques cuyos headings cumplen 'niveles'
        y opcionalmente filtrados por 'titulos' (coincidencia exacta, normalizada).
      - Un Excel 'tablas.xlsx' con todas las tablas encontradas en esos bloques.
    """
    niveles_set = set(int(n) for n in (niveles or [1, 2, 3]))
    titulo_norm_set = set(base_title(t) for t in (titulos or []))

    master = _new_doc()
    composer = Composer(master)
    todas_las_tablas: List[Tuple[str, List[List[str]]]] = []

    for a in archivos:
        name = a.get("name", "archivo.docx")
        content = a.get("content", b"")
        if not content:
            continue

        bloques = _extraer_bloques(content)

        # Encuentra cada heading que pase el filtro y su rango
        for i, b in enumerate(bloques):
            if b[0] != "h":
                continue
            lvl, txt = int(b[1]), str(b[2] or "")
            if lvl not in niveles_set:
                continue
            if enforce_whitelist and not allowed_by_whitelist(lvl, txt):
                continue
            if titulo_norm_set and base_title(txt) not in titulo_norm_set:
                continue

            # Determinar fin: siguiente heading de nivel <= actual
            j = len(bloques)
            for k in range(i + 1, len(bloques)):
                if bloques[k][0] == "h" and int(bloques[k][1]) <= lvl:
                    j = k
                    break

            # Añadir el part al composer
            _append_part(composer, bloques, i, j)

            # Recolectar tablas del rango
            for k in range(i + 1, j):
                if bloques[k][0] == "t":
                    todas_las_tablas.append((name, bloques[k][2]))

    # Serializar Word
    stream = io.BytesIO()
    composer.doc.save(stream)
    stream.seek(0)
    unificado_bytes = stream.read()

    # Serializar Excel
    xlsx_bytes = _tables_to_xlsx(todas_las_tablas)

    return {
        "unificado.docx": unificado_bytes,
        "tablas.xlsx": xlsx_bytes,
    }


def procesar_grouped(archivos: List[Dict], group_level: int,
                     titulos_objetivo: List[str],
                     enforce_whitelist: bool = False) -> Dict[str, bytes]:
    """
    Genera un DOCX por cada 'título' (texto base) en el nivel indicado.
    - group_level: 1=H1, 2=H2, 3=H3
    - titulos_objetivo: si viene vacío => usa todos los títulos encontrados
    Salida: dict {"<TituloLimpio>.docx": bytes, ...}
    """
    objetivos_norm = set(base_title(t) for t in (titulos_objetivo or []))
    composers: Dict[str, Composer] = {}
    visibles: Dict[str, str] = {}

    def get_comp(key_norm: str, visible: str) -> Composer:
        if key_norm not in composers:
            tmp = _new_doc()
            composers[key_norm] = Composer(tmp)
            if visible:
                visibles.setdefault(key_norm, visible)
        return composers[key_norm]

    # Descubrir todos los títulos nivel group_level
    all_norm: set = set()
    for a in archivos:
        content = a.get("content", b"") or b""
        if not content:
            continue
        bloques = _extraer_bloques(content)
        for kind, lvl, txt in bloques:
            if kind == "h" and int(lvl) == int(group_level) and (txt or "").strip():
                if enforce_whitelist and not allowed_by_whitelist(int(lvl), txt):
                    continue
                all_norm.add(base_title(txt))
                visibles.setdefault(base_title(txt), str(txt))

    target_keys = objetivos_norm or all_norm
    if not target_keys:
        return {}

    # Para cada archivo, apendea las secciones de cada título target
    for a in archivos:
        content = a.get("content", b"") or b""
        if not content:
            continue
        bloques = _extraer_bloques(content)

        # Índices de headings
        for i, b in enumerate(bloques):
            if b[0] != "h":
                continue
            lvl, txt = int(b[1]), str(b[2] or "")
            if lvl != group_level:
                continue
            key = base_title(txt)
            if key not in target_keys:
                continue
            if enforce_whitelist and not allowed_by_whitelist(lvl, txt):
                continue

            # fin del rango: siguiente heading de nivel <= group_level
            j = len(bloques)
            for k in range(i + 1, len(bloques)):
                if bloques[k][0] == "h" and int(bloques[k][1]) <= group_level:
                    j = k
                    break

            comp = get_comp(key, visibles.get(key, txt))
            _append_part(comp, bloques, i, j)

    # Serializar cada composer
    out: Dict[str, bytes] = {}
    for key_norm, comp in composers.items():
        visible = (visibles.get(key_norm) or key_norm).strip()
        safe = re.sub(r'[\\/:*?"<>|]+', "_", visible).strip() or key_norm or "titulo"
        filename = f"{safe}.docx"
        stream = io.BytesIO()
        comp.doc.save(stream)
        stream.seek(0)
        out[filename] = stream.read()

    return out
